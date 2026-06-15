# -*- coding: utf-8 -*-
# Part of Odoo. See LICENSE file for full copyright and licensing details.

"""
Document Parser Service
=======================
Parses uploaded documents (PDF, DOCX, TXT) and extracts text content.

Supported formats:
- PDF: Uses PyPDF2 (or pypdf) for text extraction
- DOCX: Uses python-docx for text extraction
- TXT: Raw text extraction
- HTML/XML: Basic tag stripping
"""

import base64
import logging
import re

from odoo import api, fields, models, _
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

try:
    import PyPDF2
    _HAS_PYPDF2 = True
except ImportError:
    _HAS_PYPDF2 = False
    _logger.warning('PyPDF2 not available. PDF parsing will be limited.')

try:
    from docx import Document as DocxDocument
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    _logger.warning('python-docx not available. DOCX parsing will be limited.')


class DocumentParserService(models.AbstractModel):
    """Service for parsing uploaded documents."""

    _name = 'ai.document.parser.service'
    _description = 'AI Document Parser Service'
    _inherit = 'ai.base.service'

    # ---- Max file size per type ----
    MAX_FILE_SIZE = {
        'pdf': 50 * 1024 * 1024,   # 50 MB
        'docx': 30 * 1024 * 1024,  # 30 MB
        'txt': 10 * 1024 * 1024,   # 10 MB
    }

    # ---- Main API ----

    @api.model
    def parse_document(self, document):
        """
        Parse a document and extract text content.

        :param document: ai.document record
        :return: dict with parsed data:
            - raw_text: str, the extracted plain text
            - parsed_content: str, HTML-formatted content (optional)
            - page_count: int, number of pages (if applicable)
            - metadata: dict, additional metadata
        """
        self.ensure_one_document(document)

        if not document.datas:
            raise UserError(_('No file content to parse.'))

        file_data = base64.b64decode(document.datas)
        file_type = document.document_type or self._detect_type(document.datas_fname)
        filename = document.datas_fname or 'unknown'

        self.log_info(f'Parsing document: {filename} (type: {file_type}, size: {len(file_data)} bytes)')

        # Dispatch to type-specific parser
        parser_method = f'_parse_{file_type}'
        if hasattr(self, parser_method):
            result = getattr(self, parser_method)(file_data, filename)
        else:
            result = self._parse_other(file_data, filename)

        # Ensure all required keys
        result.setdefault('raw_text', '')
        result.setdefault('parsed_content', '')
        result.setdefault('page_count', 0)
        result.setdefault('metadata', {})

        self.log_info(f'Parsing complete: {len(result["raw_text"])} chars, {result["page_count"]} pages')

        return result

    # ---- Format-Specific Parsers ----

    @api.model
    def _parse_pdf(self, file_data, filename):
        """Parse PDF document using PyPDF2."""
        if not _HAS_PYPDF2:
            return self._parse_pdf_fallback(file_data, filename)

        import io
        text_content = []
        page_count = 0

        try:
            pdf_file = io.BytesIO(file_data)
            reader = PyPDF2.PdfReader(pdf_file)
            page_count = len(reader.pages)

            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ''
                text_content.append(f'[Page {i + 1}]\n{page_text}')

            raw_text = '\n\n'.join(text_content)
            # Clean up common PDF artifacts
            raw_text = self._clean_text(raw_text)

            return {
                'raw_text': raw_text,
                'parsed_content': self._text_to_html(raw_text),
                'page_count': page_count,
                'metadata': {
                    'format': 'PDF',
                    'pages': page_count,
                    'parser': 'PyPDF2',
                },
            }

        except Exception as e:
            self.log_error(f'PDF parsing failed: {e}', exc_info=True)
            return self._parse_pdf_fallback(file_data, filename)

    @api.model
    def _parse_pdf_fallback(self, file_data, filename):
        """Fallback PDF parser using basic text extraction."""
        try:
            raw_text = file_data.decode('latin-1')
            # Try to extract readable text from PDF binary
            # This is a basic heuristic
            text_parts = re.findall(r'\((.*?)\)', raw_text)
            cleaned = '\n'.join(t for t in text_parts if len(t) > 3 and t.isprintable())
            return {
                'raw_text': cleaned or 'PDF text extraction not available. Please install PyPDF2.',
                'parsed_content': f'<pre>{cleaned}</pre>' if cleaned else '<p>PDF parsing unavailable</p>',
                'page_count': 0,
                'metadata': {'format': 'PDF', 'parser': 'fallback', 'note': 'Install PyPDF2 for better parsing'},
            }
        except Exception as e:
            return {
                'raw_text': f'[PDF parsing error: {e}]',
                'parsed_content': f'<p>Error parsing PDF: {e}</p>',
                'page_count': 0,
                'metadata': {'format': 'PDF', 'error': str(e)},
            }

    @api.model
    def _parse_docx(self, file_data, filename):
        """Parse DOCX document using python-docx."""
        if not _HAS_DOCX:
            return self._parse_docx_fallback(file_data, filename)

        import io
        try:
            docx_file = io.BytesIO(file_data)
            doc = DocxDocument(docx_file)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            raw_text = '\n\n'.join(paragraphs)

            # Extract tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text for cell in row.cells)
                    tables_text.append(row_text)

            if tables_text:
                raw_text += '\n\n=== Tables ===\n' + '\n'.join(tables_text)

            return {
                'raw_text': raw_text,
                'parsed_content': self._text_to_html(raw_text),
                'page_count': 0,  # DOCX doesn't easily expose page count
                'metadata': {
                    'format': 'DOCX',
                    'paragraphs': len(paragraphs),
                    'parser': 'python-docx',
                },
            }

        except Exception as e:
            self.log_error(f'DOCX parsing failed: {e}', exc_info=True)
            return self._parse_docx_fallback(file_data, filename)

    @api.model
    def _parse_docx_fallback(self, file_data, filename):
        """Fallback for DOCX (which is a ZIP of XML files)."""
        try:
            import zipfile
            import io
            import xml.etree.ElementTree as ET

            docx_file = io.BytesIO(file_data)
            with zipfile.ZipFile(docx_file) as z:
                # Extract the main document XML
                if 'word/document.xml' in z.namelist():
                    xml_content = z.read('word/document.xml')
                    root = ET.fromstring(xml_content)
                    # Namespace for WordprocessingML
                    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    texts = []
                    for t in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            texts.append(t.text)
                    raw_text = ' '.join(texts)
                    return {
                        'raw_text': raw_text,
                        'parsed_content': self._text_to_html(raw_text),
                        'page_count': 0,
                        'metadata': {'format': 'DOCX', 'parser': 'xml-fallback'},
                    }
        except Exception:
            pass

        return {
            'raw_text': 'DOCX parsing failed. Please install python-docx.',
            'parsed_content': '<p>DOCX parsing unavailable</p>',
            'page_count': 0,
            'metadata': {'format': 'DOCX', 'error': 'Parser unavailable'},
        }

    @api.model
    def _parse_txt(self, file_data, filename):
        """Parse plain text document."""
        try:
            # Try UTF-8 first, then latin-1
            try:
                raw_text = file_data.decode('utf-8')
            except UnicodeDecodeError:
                raw_text = file_data.decode('latin-1')

            return {
                'raw_text': raw_text,
                'parsed_content': self._text_to_html(raw_text),
                'page_count': 1,
                'metadata': {'format': 'TXT', 'size': len(raw_text)},
            }

        except Exception as e:
            return {
                'raw_text': f'[Error decoding text: {e}]',
                'parsed_content': f'<p>Error: {e}</p>',
                'page_count': 1,
                'metadata': {'format': 'TXT', 'error': str(e)},
            }

    @api.model
    def _parse_html(self, file_data, filename):
        """Parse HTML document, stripping tags."""
        try:
            raw_text = file_data.decode('utf-8')
            # Simple tag stripping
            clean_text = re.sub(r'<[^>]+>', ' ', raw_text)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()

            return {
                'raw_text': clean_text,
                'parsed_content': raw_text,  # Keep original HTML
                'page_count': 1,
                'metadata': {'format': 'HTML'},
            }
        except Exception as e:
            return self._parse_txt(file_data, filename)

    @api.model
    def _parse_markdown(self, file_data, filename):
        """Parse Markdown document."""
        try:
            raw_text = file_data.decode('utf-8')
            return {
                'raw_text': raw_text,
                'parsed_content': f'<pre>{raw_text}</pre>',
                'page_count': 1,
                'metadata': {'format': 'Markdown'},
            }
        except Exception as e:
            return self._parse_txt(file_data, filename)

    @api.model
    def _parse_csv(self, file_data, filename):
        """Parse CSV document."""
        try:
            raw_text = file_data.decode('utf-8')
            return {
                'raw_text': raw_text,
                'parsed_content': f'<pre>{raw_text}</pre>',
                'page_count': 1,
                'metadata': {'format': 'CSV', 'lines': len(raw_text.splitlines())},
            }
        except Exception as e:
            return self._parse_txt(file_data, filename)

    @api.model
    def _parse_other(self, file_data, filename):
        """Fallback for unknown file types."""
        try:
            text = file_data.decode('utf-8', errors='replace')
            if len(text) > 0 and text.isprintable():
                return {
                    'raw_text': text,
                    'parsed_content': f'<pre>{text}</pre>',
                    'page_count': 1,
                    'metadata': {'format': 'unknown', 'note': 'Parsed as raw text'},
                }
        except Exception:
            pass

        return {
            'raw_text': f'[Cannot parse file: {filename}. Unsupported format.]',
            'parsed_content': f'<p>Unsupported file format: {filename}</p>',
            'page_count': 0,
            'metadata': {'format': 'unknown', 'error': 'Unsupported format'},
        }

    # ---- Helper Methods ----

    @api.model
    def ensure_one_document(self, document):
        """Ensure the document record is valid."""
        if not document:
            raise UserError(_('No document provided.'))
        if not document.datas:
            raise UserError(_('Document has no file content.'))

    @api.model
    def _detect_type(self, filename):
        """Detect document type from filename extension."""
        if not filename:
            return 'txt'
        ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
        mapping = {
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',
            'txt': 'txt',
            'text': 'txt',
            'html': 'html',
            'htm': 'html',
            'md': 'markdown',
            'markdown': 'markdown',
            'csv': 'csv',
            'xml': 'html',
            'json': 'txt',
            'yaml': 'txt',
            'yml': 'txt',
        }
        return mapping.get(ext, 'other')

    @api.model
    def _clean_text(self, text):
        """Clean up extracted text by removing common artifacts."""
        if not text:
            return ''

        # Remove multiple blank lines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove null bytes
        text = text.replace('\x00', '')
        # Normalize unicode
        import unicodedata
        text = unicodedata.normalize('NFKC', text)
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        return text.strip()

    @api.model
    def _text_to_html(self, text):
        """Convert plain text to simple HTML."""
        if not text:
            return ''
        # Escape HTML entities
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        # Convert double newlines to paragraphs
        paragraphs = re.split(r'\n\s*\n', text)
        html_parts = [f'<p>{p.replace(chr(10), "<br/>")}</p>' for p in paragraphs if p.strip()]
        return '\n'.join(html_parts)
